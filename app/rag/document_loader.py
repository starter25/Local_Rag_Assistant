import csv
import re
import zipfile
from html import unescape
from pathlib import Path
from xml.etree import ElementTree

from docx import Document
from pypdf import PdfReader

from app.config import OCR_MIN_TEXT_CHARS
from app.rag.ocr_service import ocr_image_file, ocr_pdf_file


CODE_EXTENSIONS = {
    ".py",
    ".js",
    ".ts",
    ".tsx",
    ".jsx",
    ".html",
    ".css",
    ".json",
    ".yaml",
    ".yml",
    ".java",
    ".cpp",
    ".c",
    ".cs",
    ".go",
    ".rs",
    ".php",
    ".sql",
    ".sh",
    ".ps1",
    ".bat",
}
BASE_DOCUMENT_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}
TABLE_EXTENSIONS = {".csv", ".xlsx"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}
SUPPORTED_DOCUMENT_EXTENSIONS = (
    BASE_DOCUMENT_EXTENSIONS | CODE_EXTENSIONS | TABLE_EXTENSIONS | IMAGE_EXTENSIONS
)
MAX_TABLE_ROWS = 2000
MAX_TABLE_COLUMNS = 80
MAX_CELL_CHARS = 300

CODE_LANGUAGE_BY_EXTENSION = {
    ".py": "Python",
    ".js": "JavaScript",
    ".ts": "TypeScript",
    ".tsx": "TypeScript React",
    ".jsx": "JavaScript React",
    ".html": "HTML",
    ".css": "CSS",
    ".json": "JSON",
    ".yaml": "YAML",
    ".yml": "YAML",
    ".java": "Java",
    ".cpp": "C++",
    ".c": "C",
    ".cs": "C#",
    ".go": "Go",
    ".rs": "Rust",
    ".php": "PHP",
    ".sql": "SQL",
    ".sh": "Shell",
    ".ps1": "PowerShell",
    ".bat": "Batch",
}


# 텍스트 계열 파일은 UTF-8 우선, Windows 한글 파일을 위해 cp949까지 순서대로 시도합니다.
def read_text_with_fallback(path: Path) -> str:
    for encoding in ("utf-8", "utf-8-sig", "cp949"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue

    return path.read_text(encoding="utf-8", errors="replace")


# TXT/MD는 원문 텍스트를 그대로 RAG 입력으로 사용합니다.
def read_txt_or_md(path: Path):
    text = read_text_with_fallback(path)
    return [{"text": text, "page": None}]


# 코드 파일은 언어와 파일명을 앞에 붙여 검색 결과에서 맥락을 잃지 않게 합니다.
def read_code(path: Path):
    suffix = path.suffix.lower()
    language = CODE_LANGUAGE_BY_EXTENSION.get(suffix, suffix.lstrip(".").upper())
    text = read_text_with_fallback(path)
    wrapped = f"""File: {path.name}
Type: code
Language: {language}

{text}
""".strip()

    return [{"text": wrapped, "page": None}]


# 텍스트 PDF는 페이지 단위로 추출해 출처 표시에 page 정보를 남깁니다.
def collect_page_warnings(pages: list[dict]) -> list[str]:
    warnings = []

    for page in pages:
        for warning in page.get("warnings", []) or []:
            if warning and warning not in warnings:
                warnings.append(warning)

    return warnings


def total_text_chars(pages: list[dict]) -> int:
    return sum(len((page.get("text") or "").strip()) for page in pages)


def read_pdf(path: Path, progress=None):
    reader = PdfReader(str(path))
    pages = []

    for page_idx, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""

        if text.strip():
            pages.append(
                {
                    "text": text,
                    "page": page_idx,
                }
            )

    if total_text_chars(pages) >= OCR_MIN_TEXT_CHARS:
        return pages

    if progress:
        progress("ocr_pdf", "PDF text is sparse. Trying OCR fallback.")

    ocr_pages = ocr_pdf_file(path, progress=progress)
    warnings = collect_page_warnings(ocr_pages)

    if any((page.get("text") or "").strip() for page in ocr_pages):
        if not warnings:
            warnings.append("OCR was used because the PDF had little extractable text.")

        for page in ocr_pages:
            page["warnings"] = list(dict.fromkeys((page.get("warnings") or []) + warnings))

        return ocr_pages

    if pages:
        for page in pages:
            page["warnings"] = list(
                dict.fromkeys((page.get("warnings") or []) + warnings + ["OCR found no extra text."])
            )

        return pages

    return ocr_pages


def read_image_with_ocr(path: Path):
    result = ocr_image_file(path)
    warnings = result.get("warnings") or []

    if result.get("ocr_used") and not warnings:
        warnings = ["OCR was used to read this image."]

    return [
        {
            **result,
            "page": None,
            "warnings": warnings,
        }
    ]


# DOCX는 문단과 표 셀을 함께 텍스트로 풀어 RAG에 넣습니다.
def read_docx(path: Path):
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]

            if cells:
                parts.append(" | ".join(cells))

    text = "\n".join(parts)

    return [{"text": text, "page": None}]


# 표 셀 값은 너무 긴 값과 줄바꿈을 정리해 chunk 품질을 안정화합니다.
def clean_cell(value) -> str:
    text = str(value or "").strip()
    text = re.sub(r"\s+", " ", text)
    return text[:MAX_CELL_CHARS]


# CSV/XLSX 행을 "컬럼=값" 형태로 바꿔 LLM이 표 구조를 이해하기 쉽게 합니다.
def row_to_text(row_index: int, headers: list[str], values: list[str]) -> str:
    pairs = []

    for col_index, value in enumerate(values[:MAX_TABLE_COLUMNS]):
        value = clean_cell(value)

        if not value:
            continue

        header = clean_cell(headers[col_index]) if col_index < len(headers) else ""
        label = header or f"Column {col_index + 1}"
        pairs.append(f"{label}={value}")

    if not pairs:
        return ""

    return f"Row {row_index}: " + ", ".join(pairs)


# CSV는 첫 행을 헤더로 보고 이후 행을 검색 가능한 텍스트 행으로 변환합니다.
def read_csv_file(path: Path):
    text = read_text_with_fallback(path)
    rows = []

    for dialect in ("excel",):
        try:
            rows = list(csv.reader(text.splitlines(), dialect=dialect))
            break
        except csv.Error:
            rows = []

    if not rows:
        return [{"text": f"File: {path.name}\nType: table\nRows: 0", "page": None}]

    headers = [clean_cell(value) for value in rows[0][:MAX_TABLE_COLUMNS]]
    lines = [
        f"File: {path.name}",
        "Type: table",
        f"Rows: {max(0, len(rows) - 1)}",
        "Format: CSV",
        "",
    ]

    for row_index, row in enumerate(rows[1 : MAX_TABLE_ROWS + 1], start=1):
        row_text = row_to_text(row_index, headers, row)

        if row_text:
            lines.append(row_text)

    if len(rows) - 1 > MAX_TABLE_ROWS:
        lines.append(f"... truncated after {MAX_TABLE_ROWS} rows")

    return [{"text": "\n".join(lines), "page": None}]


# XLSX shared string과 inline string에서 실제 표시 텍스트를 꺼냅니다.
def xml_text(element, namespace: dict) -> str:
    text_parts = []

    for text_node in element.findall(".//main:t", namespace):
        text_parts.append(text_node.text or "")

    return unescape("".join(text_parts))


# XLSX 내부의 sharedStrings.xml은 셀 문자열 인덱스를 실제 문자열로 바꾸는 표입니다.
def read_xlsx_shared_strings(archive: zipfile.ZipFile) -> list[str]:
    try:
        xml_data = archive.read("xl/sharedStrings.xml")
    except KeyError:
        return []

    root = ElementTree.fromstring(xml_data)
    namespace = {"main": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}

    return [xml_text(item, namespace) for item in root.findall("main:si", namespace)]


# workbook 관계 파일을 읽어 sheet XML 경로와 사람이 보는 시트명을 연결합니다.
def read_xlsx_sheet_names(archive: zipfile.ZipFile) -> dict[str, str]:
    try:
        workbook_xml = archive.read("xl/workbook.xml")
        rels_xml = archive.read("xl/_rels/workbook.xml.rels")
    except KeyError:
        return {}

    main_ns = {"main": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
    rel_ns = {"rel": "http://schemas.openxmlformats.org/package/2006/relationships"}
    workbook_root = ElementTree.fromstring(workbook_xml)
    rels_root = ElementTree.fromstring(rels_xml)
    relationship_targets = {}

    for rel in rels_root.findall("rel:Relationship", rel_ns):
        rel_id = rel.attrib.get("Id")
        target = rel.attrib.get("Target", "")

        if rel_id and target:
            relationship_targets[rel_id] = f"xl/{target.lstrip('/')}"

    sheet_names = {}

    for sheet in workbook_root.findall(".//main:sheet", main_ns):
        name = sheet.attrib.get("name", "Sheet")
        rel_id = sheet.attrib.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
        target = relationship_targets.get(rel_id)

        if target:
            sheet_names[target] = name

    return sheet_names


# A1, C3 같은 셀 주소에서 열 번호를 계산해 빈 셀 위치도 보존합니다.
def cell_reference_to_index(reference: str) -> int:
    letters = "".join(ch for ch in reference if ch.isalpha()).upper()
    index = 0

    for char in letters:
        index = index * 26 + (ord(char) - ord("A") + 1)

    return max(0, index - 1)


# XLSX 셀 타입에 따라 shared string, inline string, 숫자 값을 공통 문자열로 읽습니다.
def read_xlsx_cell_value(cell, shared_strings: list[str], namespace: dict) -> str:
    cell_type = cell.attrib.get("t")

    if cell_type == "inlineStr":
        return xml_text(cell, namespace)

    value_node = cell.find("main:v", namespace)

    if value_node is None or value_node.text is None:
        return ""

    raw = value_node.text

    if cell_type == "s":
        try:
            return shared_strings[int(raw)]
        except (ValueError, IndexError):
            return ""

    return raw


# XLSX 한 시트를 행 텍스트 묶음으로 변환합니다.
def read_xlsx_sheet(archive: zipfile.ZipFile, sheet_path: str, sheet_name: str, shared_strings: list[str]):
    namespace = {"main": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
    root = ElementTree.fromstring(archive.read(sheet_path))
    rows = []

    for row in root.findall(".//main:sheetData/main:row", namespace):
        values = []

        for cell in row.findall("main:c", namespace):
            cell_ref = cell.attrib.get("r", "")
            col_index = cell_reference_to_index(cell_ref)

            while len(values) <= col_index:
                values.append("")

            values[col_index] = read_xlsx_cell_value(cell, shared_strings, namespace)

        if any(clean_cell(value) for value in values):
            rows.append(values[:MAX_TABLE_COLUMNS])

        if len(rows) > MAX_TABLE_ROWS:
            break

    if not rows:
        return {
            "text": f"File: {Path(sheet_path).name}\nType: spreadsheet\nSheet: {sheet_name}\nRows: 0",
            "page": sheet_name,
        }

    headers = [clean_cell(value) for value in rows[0]]
    lines = [
        f"Type: spreadsheet",
        f"Sheet: {sheet_name}",
        f"Rows: {max(0, len(rows) - 1)}",
        "",
    ]

    for row_index, row in enumerate(rows[1:], start=1):
        row_text = row_to_text(row_index, headers, row)

        if row_text:
            lines.append(row_text)

    if len(rows) > MAX_TABLE_ROWS:
        lines.append(f"... truncated after {MAX_TABLE_ROWS} rows")

    return {
        "text": "\n".join(lines),
        "page": sheet_name,
    }


# XLSX 파일은 시트별 page처럼 나누어 저장해 출처에서 시트명을 볼 수 있게 합니다.
def read_xlsx_file(path: Path):
    pages = []

    with zipfile.ZipFile(path) as archive:
        shared_strings = read_xlsx_shared_strings(archive)
        sheet_names = read_xlsx_sheet_names(archive)
        sheet_paths = sorted(
            name
            for name in archive.namelist()
            if name.startswith("xl/worksheets/sheet") and name.endswith(".xml")
        )

        for sheet_path in sheet_paths:
            sheet_name = sheet_names.get(sheet_path, Path(sheet_path).stem)
            page = read_xlsx_sheet(archive, sheet_path, sheet_name, shared_strings)
            page["text"] = f"File: {path.name}\n{page['text']}"
            pages.append(page)

    return pages


# 업로드 파일 확장자에 맞는 로더를 선택하는 단일 진입점입니다.
def load_document(path: Path, progress=None):
    suffix = path.suffix.lower()

    if suffix in [".txt", ".md"]:
        return read_txt_or_md(path)

    if suffix in CODE_EXTENSIONS:
        return read_code(path)

    if suffix == ".csv":
        return read_csv_file(path)

    if suffix == ".xlsx":
        return read_xlsx_file(path)

    if suffix == ".pdf":
        return read_pdf(path, progress=progress)

    if suffix == ".docx":
        return read_docx(path)

    if suffix in IMAGE_EXTENSIONS:
        return read_image_with_ocr(path)

    return []
